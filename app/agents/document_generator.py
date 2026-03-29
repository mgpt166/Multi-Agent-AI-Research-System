"""
app/agents/document_generator.py
==================================
DocumentGenerator — converts the annotated research narrative into a
professionally formatted Word document (.docx) using python-docx.

This is a deterministic formatter — it contains no LLM calls. It takes the
already-synthesized and citation-annotated text and applies structure,
typography, and layout to produce the final deliverable.

Report structure (5 sections, per IDEAS.md):
    1. Executive Summary        — first section of the narrative
    2. Research Findings        — organized by sub-topic with inline citations [N]
    3. Recommendations          — actionable conclusions
    4. Limitations & Gaps       — what was not found or is uncertain
    5. References / Bibliography— numbered list matching the inline [N] citations

Formatting conventions:
    - Section headings: navy blue (RGB 1F497D), bold
    - Citation markers [N]: blue (RGB 0056B8), bold, smaller font (9pt)
    - Bold text in narrative (**text**): rendered as bold runs
    - Bullet points (- item): rendered as List Bullet style
    - Markdown headings (##, ###): mapped to Word Heading 2 / Heading 3
    - Force-proceed caveat: orange italic note when synthesis review cap was hit

Input:
    query                   str         — original research question (used as subtitle)
    annotated_narrative     str         — synthesis text with [N] citation markers
    bibliography            list[dict]  — [{number, url, title, date}]
    synthesis_review_count  int         — used to determine if force-proceed caveat needed
    synthesis_review_signal str         — "force_proceed" triggers the caveat paragraph
    metadata                dict        — {generated, sources, depth, iterations}
    output_path             str         — full file path to save the .docx

Output:
    Saves the .docx file to output_path. Returns None.
    The caller (generate_document node) is responsible for creating the directory.
"""
from __future__ import annotations
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocumentGenerator:

    def generate(
        self,
        query: str,
        annotated_narrative: str,
        bibliography: list[dict],
        synthesis_review_count: int,
        synthesis_review_signal: str,
        metadata: dict,
        output_path: str,
    ) -> None:
        """Generate a clean .docx report and save to output_path."""
        doc = Document()
        self._set_document_styles(doc)

        # ── Title ──────────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"Research Report")
        run.font.size = Pt(20)
        run.font.bold = True

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_para.add_run(query[:120] + ("..." if len(query) > 120 else ""))
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True

        # Metadata line
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = (
            f"Generated: {metadata.get('generated', '')}  |  "
            f"Sources: {metadata.get('sources', 0)}  |  "
            f"Depth: {metadata.get('depth', 'moderate').title()}  |  "
            f"Iterations: {metadata.get('iterations', 1)}"
        )
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()  # spacer

        # ── Parse narrative into sections ──────────────────────────────
        sections = self._parse_narrative_sections(annotated_narrative)

        # ── 1. Executive Summary ───────────────────────────────────────
        self._add_heading(doc, "1. Executive Summary", level=1)
        exec_summary = sections.get("executive_summary", "")
        if exec_summary:
            self._add_body_text(doc, exec_summary)
        else:
            # Extract first few paragraphs as executive summary
            first_paras = "\n\n".join(annotated_narrative.split("\n\n")[:3])
            self._add_body_text(doc, first_paras or "Research in progress.")

        # ── 2. Research Findings ───────────────────────────────────────
        self._add_heading(doc, "2. Research Findings", level=1)
        findings_content = sections.get("research_findings", sections.get("findings", ""))
        if findings_content:
            self._add_narrative_with_citations(doc, findings_content)
        else:
            # Fall back to the full narrative
            self._add_narrative_with_citations(doc, annotated_narrative)

        # ── 3. Recommendations ─────────────────────────────────────────
        self._add_heading(doc, "3. Recommendations", level=1)
        recs = sections.get("recommendations", "")
        if recs:
            self._add_narrative_with_citations(doc, recs)
        else:
            doc.add_paragraph("Recommendations are derived from the research findings above.")

        # ── 4. Limitations & Gaps ──────────────────────────────────────
        self._add_heading(doc, "4. Limitations & Gaps", level=1)
        limitations = sections.get("limitations_gaps", sections.get("limitations", ""))
        if limitations:
            self._add_narrative_with_citations(doc, limitations)
        else:
            doc.add_paragraph("No major limitations identified.")

        # Force-proceed caveat
        if synthesis_review_signal == "force_proceed" and synthesis_review_count >= 3:
            caveat = doc.add_paragraph()
            caveat_run = caveat.add_run(
                "Note: This report reached the maximum synthesis review rounds. "
                "Some sub-topics may have limited coverage."
            )
            caveat_run.font.italic = True
            caveat_run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

        # ── 5. References / Bibliography ───────────────────────────────
        if bibliography:
            self._add_heading(doc, "5. References / Bibliography", level=1)
            for entry in sorted(bibliography, key=lambda x: x.get("number", 0)):
                n = entry.get("number", "?")
                title = entry.get("title", entry.get("url", "Unknown"))
                url = entry.get("url", "")
                date = entry.get("date", "")
                date_str = f". Published: {date}" if date else ""
                ref_para = doc.add_paragraph(style="List Number")
                ref_para.clear()
                # Manual formatting for reference entries
                ref_run = ref_para.add_run(f"[{n}] {title}{date_str}. ")
                ref_run.font.size = Pt(10)
                url_run = ref_para.add_run(url)
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB8)

        doc.save(output_path)

    # ── Private helpers ───────────────────────────────────────────────────────

    def _set_document_styles(self, doc: Document) -> None:
        """Set default styles for the document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    def _add_body_text(self, doc: Document, text: str) -> None:
        """Add body text, splitting on double newlines into separate paragraphs."""
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            p = doc.add_paragraph()
            # Handle bold **text**
            parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)

    def _add_narrative_with_citations(self, doc: Document, text: str) -> None:
        """
        Add narrative text, parsing markdown-style headings (##, ###) and citation markers [N].
        """
        lines = text.split("\n")
        current_para_lines: list[str] = []

        def flush_para():
            if current_para_lines:
                content = " ".join(current_para_lines).strip()
                if content:
                    self._add_inline_citation_para(doc, content)
                current_para_lines.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_para()
                continue

            # Detect headings
            if stripped.startswith("### "):
                flush_para()
                self._add_heading(doc, stripped[4:], level=3)
            elif stripped.startswith("## "):
                flush_para()
                self._add_heading(doc, stripped[3:], level=2)
            elif stripped.startswith("# "):
                flush_para()
                self._add_heading(doc, stripped[2:], level=2)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                flush_para()
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.clear()
                self._add_inline_runs(bullet, stripped[2:])
            else:
                current_para_lines.append(stripped)

        flush_para()

    def _add_inline_citation_para(self, doc: Document, text: str) -> None:
        """Add a paragraph with inline citation [N] markers rendered as superscript-style."""
        p = doc.add_paragraph()
        self._add_inline_runs(p, text)

    def _add_inline_runs(self, para, text: str) -> None:
        """Split text on [N] citation markers and bold **text**, add as runs."""
        # Split on [N] markers and **bold** markers
        tokens = re.split(r'(\[\d+\]|\*\*[^*]+\*\*)', text)
        for token in tokens:
            if re.match(r'^\[\d+\]$', token):
                run = para.add_run(token)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xB8)
                run.font.bold = True
            elif token.startswith("**") and token.endswith("**"):
                run = para.add_run(token[2:-2])
                run.font.bold = True
            else:
                para.add_run(token)

    def _parse_narrative_sections(self, narrative: str) -> dict[str, str]:
        """
        Parse the narrative into sections based on headings.
        Returns a dict of section_name -> content.
        """
        sections: dict[str, str] = {}
        current_key = "preamble"
        current_lines: list[str] = []

        section_map = {
            "executive summary": "executive_summary",
            "research findings": "research_findings",
            "findings": "research_findings",
            "recommendations": "recommendations",
            "limitations": "limitations_gaps",
            "limitations & gaps": "limitations_gaps",
            "gaps": "limitations_gaps",
        }

        for line in narrative.split("\n"):
            stripped = line.strip()
            lower = stripped.lstrip("#").strip().lower()
            matched_key = None
            for pattern, key in section_map.items():
                if lower.startswith(pattern):
                    matched_key = key
                    break

            if matched_key and (stripped.startswith("#") or stripped.startswith("**")):
                sections[current_key] = "\n".join(current_lines).strip()
                current_key = matched_key
                current_lines = []
            else:
                current_lines.append(line)

        sections[current_key] = "\n".join(current_lines).strip()
        return sections
